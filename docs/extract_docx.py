#!/usr/bin/env python3
"""Extract ALL text content from a DOCX file"""

import sys
from pathlib import Path

# Try to import python-docx, install if needed
try:
    from docx import Document
except ImportError:
    import subprocess
    print("Installing python-docx...", file=sys.stderr)
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
    from docx import Document

def extract_docx_content(docx_path):
    """Extract all paragraphs and table content from DOCX file"""
    
    doc = Document(docx_path)
    
    # Extract all paragraphs
    for para in doc.paragraphs:
        if para.text.strip():  # Only print non-empty paragraphs
            print(para.text)
    
    # Extract all tables
    for table in doc.tables:
        print("\n--- TABLE ---")
        for row in table.rows:
            cells_text = []
            for cell in row.cells:
                # Extract text from cell (may contain multiple paragraphs)
                cell_text = " ".join([p.text for p in cell.paragraphs if p.text.strip()])
                cells_text.append(cell_text)
            print(" | ".join(cells_text))
        print("--- END TABLE ---\n")

if __name__ == "__main__":
    docx_file = r"c:\Users\mnaco\Github\tontApp\docs\TontinesApp_UserStories_MVP.docx"
    
    if not Path(docx_file).exists():
        print(f"Error: File not found: {docx_file}", file=sys.stderr)
        sys.exit(1)
    
    extract_docx_content(docx_file)
